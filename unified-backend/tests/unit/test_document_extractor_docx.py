"""Unit tests for DOCX extraction in the document_extractor module."""

import os
import tempfile
import zipfile

import pytest

from app.services.document_extractor import extract_docx


def _create_test_docx(paragraphs: list[tuple[str, str]]) -> str:
    """Create a minimal DOCX file with specified paragraphs.

    Args:
        paragraphs: List of (style_name, text) tuples.
            style_name can be: "Normal", "Heading 1", "Heading 2",
            "List Bullet", "List Paragraph", etc.

    Returns:
        Path to the temporary DOCX file.
    """
    import docx

    doc = docx.Document()
    for style_name, text in paragraphs:
        p = doc.add_paragraph(text)
        try:
            p.style = doc.styles[style_name]
        except KeyError:
            # If style doesn't exist, use Normal
            p.style = doc.styles["Normal"]

    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(tmp.name)
    tmp.close()
    return tmp.name


def _create_corrupted_docx() -> str:
    """Create a file that looks like a DOCX but is corrupted."""
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    tmp.write(b"This is not a valid DOCX file at all")
    tmp.close()
    return tmp.name


def _create_invalid_zip_docx() -> str:
    """Create a valid ZIP that doesn't have DOCX XML structure."""
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    with zipfile.ZipFile(tmp.name, "w") as zf:
        zf.writestr("random.txt", "This is just a zip file, not a DOCX")
    return tmp.name


class TestExtractDocx:
    """Tests for extract_docx function."""

    @pytest.mark.asyncio
    async def test_extracts_normal_paragraphs(self):
        """Should extract plain paragraphs separated by double newlines."""
        filepath = _create_test_docx([
            ("Normal", "First paragraph."),
            ("Normal", "Second paragraph."),
        ])
        try:
            result = await extract_docx(filepath)
            assert "First paragraph." in result
            assert "Second paragraph." in result
            assert "\n\n" in result
        finally:
            os.unlink(filepath)

    @pytest.mark.asyncio
    async def test_extracts_headings(self):
        """Should extract heading text from the document."""
        filepath = _create_test_docx([
            ("Heading 1", "Main Title"),
            ("Normal", "Some content below the heading."),
            ("Heading 2", "Sub Section"),
            ("Normal", "More content here."),
        ])
        try:
            result = await extract_docx(filepath)
            assert "Main Title" in result
            assert "Sub Section" in result
            assert "Some content below the heading." in result
            assert "More content here." in result
        finally:
            os.unlink(filepath)

    @pytest.mark.asyncio
    async def test_extracts_list_items_with_bullet_markers(self):
        """Should prefix list items with bullet markers."""
        filepath = _create_test_docx([
            ("Normal", "Here is a list:"),
            ("List Bullet", "Item one"),
            ("List Bullet", "Item two"),
            ("List Bullet", "Item three"),
        ])
        try:
            result = await extract_docx(filepath)
            assert "• Item one" in result
            assert "• Item two" in result
            assert "• Item three" in result
        finally:
            os.unlink(filepath)

    @pytest.mark.asyncio
    async def test_extracts_list_paragraph_style(self):
        """Should prefix List Paragraph styled items with bullet markers."""
        filepath = _create_test_docx([
            ("Normal", "Introduction"),
            ("List Paragraph", "First list item"),
            ("List Paragraph", "Second list item"),
        ])
        try:
            result = await extract_docx(filepath)
            assert "• First list item" in result
            assert "• Second list item" in result
        finally:
            os.unlink(filepath)

    @pytest.mark.asyncio
    async def test_skips_empty_paragraphs(self):
        """Should skip paragraphs with only whitespace."""
        filepath = _create_test_docx([
            ("Normal", "Before empty"),
            ("Normal", ""),
            ("Normal", "   "),
            ("Normal", "After empty"),
        ])
        try:
            result = await extract_docx(filepath)
            assert "Before empty" in result
            assert "After empty" in result
            # The empty paragraphs should not produce extra separators
            parts = [p for p in result.split("\n\n") if p.strip()]
            assert len(parts) == 2
        finally:
            os.unlink(filepath)

    @pytest.mark.asyncio
    async def test_discards_formatting_metadata(self):
        """Should only extract text, no font/color/spacing info."""
        import docx
        from docx.shared import Pt, RGBColor

        doc = docx.Document()
        p = doc.add_paragraph()
        run = p.add_run("Formatted text")
        run.bold = True
        run.italic = True
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(255, 0, 0)

        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        doc.save(tmp.name)
        tmp.close()

        try:
            result = await extract_docx(tmp.name)
            assert "Formatted text" in result
            # No formatting metadata in output
            assert "bold" not in result.lower()
            assert "italic" not in result.lower()
            assert "rgb" not in result.lower()
            assert "255" not in result
        finally:
            os.unlink(tmp.name)

    @pytest.mark.asyncio
    async def test_raises_on_corrupted_file(self):
        """Should raise ValueError for corrupted/non-DOCX files."""
        filepath = _create_corrupted_docx()
        try:
            with pytest.raises(ValueError, match="corrupted or unreadable file"):
                await extract_docx(filepath)
        finally:
            os.unlink(filepath)

    @pytest.mark.asyncio
    async def test_raises_on_invalid_zip_structure(self):
        """Should raise ValueError for ZIP files without DOCX structure."""
        filepath = _create_invalid_zip_docx()
        try:
            with pytest.raises(ValueError):
                await extract_docx(filepath)
        finally:
            os.unlink(filepath)

    @pytest.mark.asyncio
    async def test_empty_document_returns_empty_string(self):
        """Should return empty string for a DOCX with no content."""
        filepath = _create_test_docx([])
        try:
            result = await extract_docx(filepath)
            assert result == ""
        finally:
            os.unlink(filepath)

    @pytest.mark.asyncio
    async def test_mixed_content_types(self):
        """Should handle a document with headings, paragraphs, and lists mixed."""
        filepath = _create_test_docx([
            ("Heading 1", "Resume"),
            ("Normal", "John Doe, Software Engineer"),
            ("Heading 2", "Skills"),
            ("List Bullet", "Python"),
            ("List Bullet", "FastAPI"),
            ("List Bullet", "MongoDB"),
            ("Heading 2", "Experience"),
            ("Normal", "5 years at Tech Corp"),
        ])
        try:
            result = await extract_docx(filepath)
            assert "Resume" in result
            assert "John Doe, Software Engineer" in result
            assert "Skills" in result
            assert "• Python" in result
            assert "• FastAPI" in result
            assert "• MongoDB" in result
            assert "Experience" in result
            assert "5 years at Tech Corp" in result
        finally:
            os.unlink(filepath)

    @pytest.mark.asyncio
    async def test_nonexistent_file_raises_error(self):
        """Should raise ValueError for a file path that doesn't exist."""
        with pytest.raises((ValueError, FileNotFoundError, OSError)):
            await extract_docx("/nonexistent/path/file.docx")
